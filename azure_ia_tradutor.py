# -*- coding: utf-8 -*-
"""Azure-IA-tradutor.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1cvDJ30Bvn4eqWZrtRuru7zUutKlEaD6p
"""

!pip install requests python-docx
!pip install os

import requests
import os
from docx import Document

subscription_key = "xGg08oQsylu13CSe2C2VBDkcH6kDFteGURaJlBEqalmdG1jE9SlKJQQJ99AKACHYHv6XJ3w3AAAbACOG4ZEa"
endpoint = "https://api.cognitive.microsofttranslator.com/"
location = "eastus2"
language_destination = 'pt-br'

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-Type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]

    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': language_destination
    }

    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]['translations'][0]['text']

translator_text("I try to face the fight within But it's over", language_destination)

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    print(line)
    translated_doc.add_paragraph(line)
    path_translated = path.replace('.docx', f'_{language_destination}.docx')
    translated_doc.save('path_translated')
    return path_translated

input_file = "/content/I try to face the fight within.docx"
translate_document(input_file)